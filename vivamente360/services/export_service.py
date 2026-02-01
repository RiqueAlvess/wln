from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from datetime import datetime
from bs4 import BeautifulSoup
import re

class ExportService:
    @staticmethod
    def export_plano_acao_word(campaign, planos):
        doc = Document()

        # Cabeçalho
        heading = doc.add_heading(f'Plano de Ação - {campaign.nome}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Empresa: {campaign.empresa.nome}')
        doc.add_paragraph(f'Período: {campaign.data_inicio} a {campaign.data_fim}')
        doc.add_paragraph('')

        # Tabela de Riscos
        doc.add_heading('Riscos Identificados', level=1)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Dimensão'
        hdr_cells[1].text = 'Nível de Risco'
        hdr_cells[2].text = 'Ação Proposta'
        hdr_cells[3].text = 'Responsável'
        hdr_cells[4].text = 'Prazo'

        for plano in planos:
            row_cells = table.add_row().cells
            row_cells[0].text = plano.dimensao.nome
            row_cells[1].text = plano.nivel_risco
            row_cells[2].text = plano.acao_proposta[:100] + '...'
            row_cells[3].text = plano.responsavel
            row_cells[4].text = str(plano.prazo)

        # Detalhamento
        doc.add_page_break()
        doc.add_heading('Detalhamento das Ações', level=1)

        for i, plano in enumerate(planos, 1):
            doc.add_heading(f'Ação {i}: {plano.dimensao.nome}', level=2)
            doc.add_paragraph(f'Nível de Risco: {plano.nivel_risco}')
            doc.add_paragraph(f'Descrição do Risco: {plano.descricao_risco}')
            doc.add_paragraph(f'Ação Proposta: {plano.acao_proposta}')
            doc.add_paragraph(f'Responsável: {plano.responsavel}')
            doc.add_paragraph(f'Prazo: {plano.prazo}')
            if plano.recursos_necessarios:
                doc.add_paragraph(f'Recursos: {plano.recursos_necessarios}')
            if plano.indicadores:
                doc.add_paragraph(f'Indicadores: {plano.indicadores}')
            doc.add_paragraph('')

        # Assinaturas
        doc.add_page_break()
        doc.add_heading('Aprovações', level=1)
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Responsável pela Área de SST')
        doc.add_paragraph('')
        doc.add_paragraph('_' * 40)
        doc.add_paragraph('Gerente de RH')

        return doc

    @staticmethod
    def export_plano_acao_rich_word(plano_acao):
        """
        Exporta um único plano de ação com conteúdo rico (HTML do editor TipTap) para DOCX
        """
        doc = Document()

        # Cabeçalho do documento
        heading = doc.add_heading(f'PLANO DE AÇÃO', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informações básicas
        doc.add_paragraph(f'Empresa: {plano_acao.empresa.nome}')
        doc.add_paragraph(f'Campanha: {plano_acao.campaign.nome}')
        doc.add_paragraph(f'Dimensão: {plano_acao.dimensao.nome}')
        doc.add_paragraph(f'Nível de Risco: {plano_acao.nivel_risco}')
        doc.add_paragraph(f'Responsável: {plano_acao.responsavel}')
        doc.add_paragraph(f'Prazo: {plano_acao.prazo.strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Status: {plano_acao.get_status_display()}')
        doc.add_paragraph('')

        # Conteúdo rico do editor (se disponível)
        if plano_acao.conteudo_html:
            doc.add_heading('Plano de Ação Detalhado', level=1)
            ExportService._html_to_docx(doc, plano_acao.conteudo_html)
        else:
            # Fallback para campos legados
            doc.add_heading('Descrição do Risco', level=1)
            doc.add_paragraph(plano_acao.descricao_risco)

            doc.add_heading('Ação Proposta', level=1)
            doc.add_paragraph(plano_acao.acao_proposta)

            if plano_acao.recursos_necessarios:
                doc.add_heading('Recursos Necessários', level=1)
                doc.add_paragraph(plano_acao.recursos_necessarios)

            if plano_acao.indicadores:
                doc.add_heading('Indicadores de Acompanhamento', level=1)
                doc.add_paragraph(plano_acao.indicadores)

        # Rodapé com assinaturas
        doc.add_page_break()
        doc.add_heading('Aprovações', level=1)
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Responsável pela Área de SST')
        doc.add_paragraph('')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Gerente de RH')
        doc.add_paragraph('')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Diretor/Presidente')

        return doc

    @staticmethod
    def _html_to_docx(doc, html_content):
        """
        Converte HTML do editor TipTap para elementos do python-docx
        """
        soup = BeautifulSoup(html_content, 'html.parser')

        for element in soup.children:
            if element.name is None:  # Texto puro
                continue

            # Cabeçalhos
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)

            # Parágrafos
            elif element.name == 'p':
                p = doc.add_paragraph()
                ExportService._add_formatted_text(p, element)

            # Listas não ordenadas
            elif element.name == 'ul':
                # Verificar se é task list
                if element.get('data-type') == 'taskList':
                    for li in element.find_all('li', recursive=False):
                        checkbox = li.find('input', {'type': 'checkbox'})
                        checked = checkbox.get('checked') if checkbox else False
                        text = li.get_text().strip()
                        symbol = '☑' if checked else '☐'
                        doc.add_paragraph(f'{symbol} {text}', style='List Bullet')
                else:
                    for li in element.find_all('li', recursive=False):
                        doc.add_paragraph(li.get_text(), style='List Bullet')

            # Listas ordenadas
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Number')

            # Tabelas
            elif element.name == 'table':
                ExportService._add_table_to_docx(doc, element)

            # Quebra de linha
            elif element.name == 'br':
                doc.add_paragraph('')

    @staticmethod
    def _add_formatted_text(paragraph, element):
        """
        Adiciona texto formatado (negrito, itálico, etc.) a um parágrafo
        """
        for child in element.children:
            if child.name is None:  # Texto puro
                run = paragraph.add_run(str(child))
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == 's':
                run = paragraph.add_run(child.get_text())
                run.font.strike = True
            elif child.name == 'span':
                run = paragraph.add_run(child.get_text())
                # Processar estilos inline (cores, etc.)
                style = child.get('style', '')
                if 'color:' in style:
                    color_match = re.search(r'color:\s*([#\w]+)', style)
                    if color_match:
                        color_hex = color_match.group(1).replace('#', '')
                        if len(color_hex) == 6:
                            try:
                                r = int(color_hex[0:2], 16)
                                g = int(color_hex[2:4], 16)
                                b = int(color_hex[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                            except ValueError:
                                pass
            else:
                # Recursão para elementos aninhados
                ExportService._add_formatted_text(paragraph, child)

    @staticmethod
    def _add_table_to_docx(doc, table_element):
        """
        Adiciona uma tabela HTML ao documento DOCX
        """
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Contar colunas
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))

        # Criar tabela
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'

        # Preencher células
        for i, tr in enumerate(rows):
            cells = tr.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                table.rows[i].cells[j].text = cell.get_text().strip()

                # Negrito para headers
                if cell.name == 'th':
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    @staticmethod
    def export_checklist_nr1_pdf(campaign, itens_por_etapa, progresso_geral, total_itens, total_concluidos):
        """
        Exporta o checklist NR-1 completo para PDF
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)

        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#0d6efd'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#212529'),
            spaceAfter=12,
            spaceBefore=12
        )

        normal_style = styles['Normal']

        # Conteúdo do documento
        story = []

        # Título
        story.append(Paragraph("Checklist de Compliance NR-1", title_style))
        story.append(Spacer(1, 0.2*inch))

        # Informações da Campanha
        info_data = [
            ['Campanha:', campaign.nome],
            ['Empresa:', campaign.empresa.nome],
            ['Período:', f'{campaign.data_inicio.strftime("%d/%m/%Y")} a {campaign.data_fim.strftime("%d/%m/%Y")}'],
            ['Data de Geração:', datetime.now().strftime("%d/%m/%Y %H:%M")],
        ]

        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e9ecef')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))

        # Progresso Geral
        progresso_data = [
            ['Progresso Geral:', f'{progresso_geral:.0f}%'],
            ['Itens Concluídos:', f'{total_concluidos} de {total_itens}'],
        ]

        progresso_table = Table(progresso_data, colWidths=[2*inch, 4*inch])
        progresso_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#d1ecf1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        story.append(progresso_table)
        story.append(Spacer(1, 0.3*inch))

        # Itens por Etapa
        for etapa_num, etapa_data in sorted(itens_por_etapa.items()):
            # Cabeçalho da Etapa
            story.append(Paragraph(f"ETAPA {etapa_num}: {etapa_data['nome'].upper()}", heading_style))
            story.append(Paragraph(
                f"Progresso: {etapa_data['progresso']:.0f}% ({etapa_data['concluidos']}/{etapa_data['total']} itens)",
                normal_style
            ))
            story.append(Spacer(1, 0.1*inch))

            # Tabela de Itens
            items_data = [['Status', 'Item', 'Responsável', 'Prazo', 'Evidências']]

            for item in etapa_data['itens']:
                status = '✓' if item.concluido else '○'
                item_texto = item.item_texto[:60] + '...' if len(item.item_texto) > 60 else item.item_texto
                if item.automatico:
                    item_texto += ' (Automático)'
                responsavel = item.responsavel if item.responsavel else '-'
                prazo = item.prazo.strftime("%d/%m/%Y") if item.prazo else '-'
                evidencias = str(item.evidencias.count())

                items_data.append([status, item_texto, responsavel, prazo, evidencias])

            items_table = Table(items_data, colWidths=[0.5*inch, 3*inch, 1.5*inch, 0.8*inch, 0.7*inch])
            items_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0d6efd')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ]))

            story.append(items_table)
            story.append(Spacer(1, 0.2*inch))

        # Assinaturas
        story.append(PageBreak())
        story.append(Paragraph("APROVAÇÕES", heading_style))
        story.append(Spacer(1, 0.3*inch))

        assinaturas_data = [
            ['_' * 50],
            ['Responsável pela Área de SST'],
            [''],
            ['_' * 50],
            ['Gerente de RH'],
            [''],
            ['_' * 50],
            ['Diretor/Presidente'],
        ]

        assinaturas_table = Table(assinaturas_data, colWidths=[6*inch])
        assinaturas_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        story.append(assinaturas_table)

        # Gerar PDF
        doc.build(story)
        pdf = buffer.getvalue()
        buffer.close()

        return pdf

    @staticmethod
    def export_campaign_comparison_word(campaign1, campaign2, summary, dimensions, sectors, ai_analysis):
        """
        Exporta relatório de comparação entre campanhas em formato Word.
        """
        doc = Document()

        # Cabeçalho
        heading = doc.add_heading(f'Relatório de Evolução - {campaign1.empresa.nome}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informações gerais
        doc.add_paragraph(f'Data de geração: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph('')

        # Campanhas comparadas
        doc.add_heading('Campanhas Comparadas', level=1)
        doc.add_paragraph(f'Campanha Base: {campaign1.nome}')
        doc.add_paragraph(f'Período: {campaign1.data_inicio} a {campaign1.data_fim}')
        doc.add_paragraph('')
        doc.add_paragraph(f'Campanha Nova: {campaign2.nome}')
        doc.add_paragraph(f'Período: {campaign2.data_inicio} a {campaign2.data_fim}')
        doc.add_paragraph('')

        # Resumo da Evolução
        doc.add_heading('Resumo da Evolução', level=1)

        table = doc.add_table(rows=5, cols=4)
        table.style = 'Light Grid Accent 1'

        # Cabeçalho da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = campaign1.nome[:20]
        hdr_cells[2].text = campaign2.nome[:20]
        hdr_cells[3].text = 'Variação'

        # Dados
        table.rows[1].cells[0].text = 'Taxa de Adesão (%)'
        table.rows[1].cells[1].text = str(summary['campaign1']['adesao'])
        table.rows[1].cells[2].text = str(summary['campaign2']['adesao'])
        variacao_adesao = summary['variacao']['adesao']
        table.rows[1].cells[3].text = f"{'+' if variacao_adesao > 0 else ''}{variacao_adesao}%"

        table.rows[2].cells[0].text = 'IGRP'
        table.rows[2].cells[1].text = str(summary['campaign1']['igrp'])
        table.rows[2].cells[2].text = str(summary['campaign2']['igrp'])
        variacao_igrp = summary['variacao']['igrp']
        table.rows[2].cells[3].text = f"{'+' if variacao_igrp > 0 else ''}{variacao_igrp}"

        table.rows[3].cells[0].text = '% Risco Alto/Crítico'
        table.rows[3].cells[1].text = str(summary['campaign1']['pct_risco_alto'])
        table.rows[3].cells[2].text = str(summary['campaign2']['pct_risco_alto'])
        variacao_risco = summary['variacao']['pct_risco_alto']
        table.rows[3].cells[3].text = f"{'+' if variacao_risco > 0 else ''}{variacao_risco}%"

        table.rows[4].cells[0].text = 'Total de Respostas'
        table.rows[4].cells[1].text = str(summary['campaign1']['total_respostas'])
        table.rows[4].cells[2].text = str(summary['campaign2']['total_respostas'])
        variacao_respostas = summary['variacao']['total_respostas']
        table.rows[4].cells[3].text = f"{'+' if variacao_respostas > 0 else ''}{variacao_respostas}"

        doc.add_paragraph('')

        # Evolução por Dimensão
        doc.add_heading('Evolução por Dimensão HSE-IT', level=1)

        dim_table = doc.add_table(rows=len(dimensions) + 1, cols=5)
        dim_table.style = 'Light Grid Accent 1'

        # Cabeçalho
        hdr_cells = dim_table.rows[0].cells
        hdr_cells[0].text = 'Dimensão'
        hdr_cells[1].text = campaign1.nome[:15]
        hdr_cells[2].text = campaign2.nome[:15]
        hdr_cells[3].text = 'Variação'
        hdr_cells[4].text = 'Tendência'

        # Dados
        for i, dim in enumerate(dimensions, 1):
            row_cells = dim_table.rows[i].cells
            row_cells[0].text = dim['dimensao']
            row_cells[1].text = str(dim['score_c1'])
            row_cells[2].text = str(dim['score_c2'])
            row_cells[3].text = f"{'+' if dim['variacao'] > 0 else ''}{dim['variacao']}"

            if dim['tendencia'] == 'melhora':
                row_cells[4].text = '↑ Melhora'
            elif dim['tendencia'] == 'piora':
                row_cells[4].text = '↓ Piora'
            else:
                row_cells[4].text = '→ Estável'

        doc.add_paragraph('')

        # Setores que Mais Melhoraram
        if sectors.get('melhoraram'):
            doc.add_heading('Setores que Mais Melhoraram', level=1)

            for i, sector in enumerate(sectors['melhoraram'], 1):
                doc.add_paragraph(
                    f"{i}. {sector['setor']}: IGRP {sector['igrp_c1']} → {sector['igrp_c2']} "
                    f"({sector['variacao']} pontos)",
                    style='List Bullet'
                )

            doc.add_paragraph('')

        # Setores que Precisam Atenção
        if sectors.get('pioraram'):
            doc.add_heading('Setores que Precisam Atenção', level=1)

            for i, sector in enumerate(sectors['pioraram'], 1):
                doc.add_paragraph(
                    f"{i}. {sector['setor']}: IGRP {sector['igrp_c1']} → {sector['igrp_c2']} "
                    f"(+{sector['variacao']} pontos)",
                    style='List Bullet'
                )

            doc.add_paragraph('')

        # Análise de Evolução (IA)
        if ai_analysis:
            doc.add_heading('Análise de Evolução (IA)', level=1)
            doc.add_paragraph(ai_analysis)
            doc.add_paragraph('')

        # Recomendações
        doc.add_page_break()
        doc.add_heading('Recomendações', level=1)

        if summary['variacao']['igrp'] < 0:
            doc.add_paragraph(
                'Manutenção: Continue implementando as ações que demonstraram eficácia na redução de riscos.',
                style='List Bullet'
            )
        else:
            doc.add_paragraph(
                'Atenção: Revise as ações implementadas e identifique oportunidades de melhoria.',
                style='List Bullet'
            )

        if sectors.get('melhoraram'):
            doc.add_paragraph(
                f'Boas práticas: Documente e replique as estratégias bem-sucedidas '
                f'implementadas no setor "{sectors["melhoraram"][0]["setor"]}".',
                style='List Bullet'
            )

        if sectors.get('pioraram'):
            doc.add_paragraph(
                f'Prioridade: Desenvolva plano de ação emergencial para o setor '
                f'"{sectors["pioraram"][0]["setor"]}".',
                style='List Bullet'
            )

        # Assinaturas
        doc.add_page_break()
        doc.add_heading('Aprovações', level=1)
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Responsável pela Área de SST')
        doc.add_paragraph('')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Gerente de RH')
        doc.add_paragraph('')
        doc.add_paragraph('_' * 50)
        doc.add_paragraph('Diretor/Presidente')

        return doc
